import os
import time
import google.generativeai as genai
from PIL import Image
import io
import edge_tts
import PyPDF2
import docx
import config
import re

# --- SETUP GOOGLE GEMINI ---
genai.configure(api_key=config.GEMINI_API_KEY)

MODEL_PRIMARY = 'gemini-2.5-flash'
MODEL_BACKUP = 'gemini-2.0-flash'

LANG_NAMES = {
    'id': 'Indonesian', 'en': 'English', 'ja': 'Japanese', 'ko': 'Korean', 'ar': 'Arabic'
}

def model_backup(content, retries=3, delay=5):
    # Tentukan model awal
    current_model_name = MODEL_PRIMARY
    
    for attempt in range(retries):
        try:
            model = genai.GenerativeModel(current_model_name)
            response = model.generate_content(content)
            return response
            
        except Exception as e:
            error_msg = str(e).lower()
            
            # Model Tidak Ditemukan (404)
            if "404" in error_msg or "not found" in error_msg:
                print(f"Model {current_model_name} tidak ditemukan (404). Menggunakan backup: {MODEL_BACKUP}")
                current_model_name = MODEL_BACKUP 
                continue 
                
            # Kena Limit Kuota (429)
            elif "429" in error_msg or "quota" in error_msg:
                print(f"⏳ Limit Kuota (429) di {current_model_name}. Menunggu {delay} detik...")
                time.sleep(delay)
                delay *= 2
                
            else:
                print(f"Error Gemini ({current_model_name}): {e}")
                if attempt == retries - 1:
                    return None
                    
    return None

# --- 1. OCR ---
def ocr_with_gemini(image_bytes) -> str:
    try:
        image = Image.open(io.BytesIO(image_bytes))
        prompt = "Extract all text from this image exactly as it appears. Do not summarize yet."
        
        response = model_backup([prompt, image])
        
        if response and response.text:
            return response.text.strip()
        return "Gagal mengekstrak teks (OCR gagal/Limit)."
    except Exception as e:
        print(f"Error Gemini OCR: {e}")
        return None

# --- 2. SUMMARIZER ---
def summarize_text(text: str) -> str:
    try:
        prompt = (
            "You are a professional editor. Your goal is to summarize the text concisely. "
            "1. The summary MUST be significantly shorter than the original text. "
            "2. Do NOT use introductory phrases like 'Here is the summary', 'Halo pendengar', or 'Berikut ringkasannya'. "
            "3. Start directly with the main content. "
            "4. Combine main ideas into flowing paragraphs. "
            "5. Do NOT use bullet points. "
            "6. Respond in the same language as the original text.\n\n"
            f"TEXT TO SUMMARIZE:\n{text[:config.MAX_CHARS]}"
        )
        
        response = model_backup(prompt)
        
        if not response or not response.text:
            return "Maaf, AI gagal merespons (Limit Kuota). Silakan coba lagi nanti."
            
        clean_text = response.text.strip()
        clean_text = clean_text.replace("* ", "").replace("- ", "").replace("• ", "").replace("**", "")
        return clean_text

    except Exception as e:
        return f"Gagal Meringkas. Error: {str(e)}"

# --- 3. TRANSLATOR ---
def translate_text(text: str, target_lang_code: str) -> str:
    try:
        target_lang_name = LANG_NAMES.get(target_lang_code, 'English')
        prompt = (
            f"Translate the following text into natural, native-sounding {target_lang_name}. "
            "STRICT RULES:\n"
            f"1. If the text is ALREADY in {target_lang_name}, RETURN IT EXACTLY AS IS. DO NOT PARAPHRASE.\n"
            "2. Translate accurately without adding explanations.\n"
            "3. Do not add introductory phrases.\n\n"
            f"TEXT:\n{text[:config.MAX_CHARS]}"
        )
        
        response = model_backup(prompt)
        
        if not response or not response.text: 
            print("Translate Gagal (Limit), mengembalikan teks asli.")
            return text 
            
        return response.text.strip()
    except Exception as e:
        print(f"Error Translate: {e}")
        return text 

# --- 4. TTS GENERATOR ---
def clean_text_for_tts(text: str) -> str:
    """Membersihkan teks dari karakter markdown agar TTS tidak bingung."""
    # Hapus bold/italic marker markdown (* atau _)
    text = text.replace("**", "").replace("__", "").replace("*", "")
    # Hapus karakter aneh yang mungkin tidak terbaca
    text = re.sub(r'[^\w\s.,?!-]', '', text) 
    return text.strip()

def split_text_smartly(text, limit):
    chunks = []
    current_chunk = ""
    if not text: return []
    
    # Bersihkan teks dulu
    text = clean_text_for_tts(text)
    
    # Split berdasarkan kalimat
    sentences = text.replace('\n', ' ').split('. ') 
    
    for sentence in sentences:
        if len(current_chunk) + len(sentence) < limit:
            current_chunk += sentence + ". "
        else:
            if current_chunk: chunks.append(current_chunk.strip())
            current_chunk = sentence + ". "
    if current_chunk:
        chunks.append(current_chunk.strip())
    return chunks

async def generate_audio_long(text: str, lang: str, gender: str, user_id: str, progress_callback=None) -> str:
    try:
        if not text or not text.strip():
            print("Error TTS: Teks input kosong.")
            return None

        voice_dict = config.VOICE_MAPPING.get(lang, config.VOICE_MAPPING['en'])
        selected_voice = voice_dict.get(gender, voice_dict['female'])
        
        chunks = split_text_smartly(text, config.CHUNK_SIZE)
        total_chunks = len(chunks)
        temp_files = []

        if total_chunks == 0:
            print("Error TTS: Tidak ada chunk teks yang dihasilkan.")
            return None

        for i, chunk in enumerate(chunks):
            if not chunk.strip(): continue
            
            # Update Progress Bar di Telegram 
            if progress_callback:
                await progress_callback(i + 1, total_chunks)

            temp_file = f"temp_{user_id}_part{i}.mp3"
            
            # RETRY LOGIC UNTUK TTS
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    communicate = edge_tts.Communicate(chunk, selected_voice)
                    await communicate.save(temp_file)
                    # Cek size file
                    if os.path.getsize(temp_file) > 0:
                        temp_files.append(temp_file)
                        break # Sukses, keluar loop retry
                    else:
                        print(f"Warning: Chunk {i} 0 bytes. Retry {attempt+1}...")
                except Exception as e:
                    print(f"Error Chunk {i} (Percobaan {attempt+1}): {e}")
                    if attempt == max_retries - 1:
                        print(f"Gagal total pada chunk {i}")
                    time.sleep(1)

        # Cek apakah ada file yang berhasil dibuat
        if not temp_files:
            print("Error TTS: Tidak ada file audio yang berhasil dibuat.")
            return None

        # Merge
        final_filename = f"audio_{user_id}_final.mp3"
        with open(final_filename, 'wb') as outfile:
            for f_path in temp_files:
                if os.path.exists(f_path):
                    with open(f_path, 'rb') as infile:
                        outfile.write(infile.read())
                    os.remove(f_path) 

        return final_filename

    except Exception as e:
        print(f"CRITICAL Error TTS Long: {e}")
        return None

# --- 5. FILE EXTRACTOR ---
def extract_document_content(file_path: str) -> str:
    text = ""
    try:
        if file_path.endswith('.pdf'):
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    t = page.extract_text()
                    if t: text += t + "\n"
        elif file_path.endswith('.docx'):
            doc = docx.Document(file_path)
            for para in doc.paragraphs: text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs: text += para.text + "\n"
        elif file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f: text = f.read()
    except Exception as e:
        print(f"Error Read File: {e}")
        return None
    return text.strip()